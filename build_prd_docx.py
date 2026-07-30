from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "用户标签物料产品需求文档_PRD.md"
OUT_PATH = ROOT / "用户标签物料产品需求文档_PRD.docx"
IMAGE_PATH = Path(
    "/Users/xinyintiaodong/Library/Containers/com.bytedance.macos.feishu/"
    "Data/Library/Application Support/LarkShell/sdk_storage/"
    "7dbceb4e66f1ae6b7ad2db50584ab24f/resources/images/"
    "img_v3_0213k_49c982a1-63d7-4f1b-839c-55f8ca2f138g.jpg"
)


INK = "1F2937"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EAF2F8"
BORDER = "CBD5E1"
WHITE = "FFFFFF"
GOLD = "A66B00"


def set_run_font(run, ascii_name="Arial", east_asia="PingFang SC", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = ascii_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), east_asia)
    rfonts.set(qn("w:hint"), "eastAsia")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_paragraph_border_bottom(paragraph, color=BLUE, size=10, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text, default_size=11, default_color=INK):
    pattern = re.compile(
        r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, ascii_name="Menlo", east_asia="PingFang SC",
                         size=max(default_size - 0.5, 8), color=DARK_BLUE)
            run.font.highlight_color = None
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            url = token.rstrip("。；，,;)")
            suffix = token[len(url):]
            add_hyperlink(paragraph, url, url)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size=default_size, color=default_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size, color=default_color)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal._element.rPr.rFonts.set(qn("w:cs"), "PingFang SC")
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style._element.rPr.rFonts.set(qn("w:cs"), "PingFang SC")
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Table Text" not in styles:
        table_style = styles.add_style("Table Text", 1)
    else:
        table_style = styles["Table Text"]
    table_style.font.name = "Arial"
    table_style.font.size = Pt(9)
    table_style.font.color.rgb = RGBColor.from_string(INK)
    table_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    table_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    table_style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    table_style._element.rPr.rFonts.set(qn("w:cs"), "PingFang SC")
    table_style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(0)
    table_style.paragraph_format.line_spacing = 1.08

    if "Caption Custom" not in styles:
        caption = styles.add_style("Caption Custom", 1)
    else:
        caption = styles["Caption Custom"]
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    caption._element.rPr.rFonts.set(qn("w:cs"), "PingFang SC")
    caption._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def create(abstract_id, num_id, fmt, text_value, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    create(next_abs, next_num, "bullet", "•", "Arial")
    create(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def new_num_instance(doc, base_num_id):
    numbering = doc.part.numbering_part.element
    base = None
    for item in numbering.findall(qn("w:num")):
        if int(item.get(qn("w:numId"))) == int(base_num_id):
            base = item
            break
    if base is None:
        return base_num_id
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    used = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    new_id = max(used, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), abstract_id)
    num.append(abs_id)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return new_id


def configure_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def populate_header(header):
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("产品需求文档  |  用户标签物料")
    set_run_font(run, size=8.5, color=MUTED, bold=True)


def populate_footer(footer):
    p = footer.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("内部评审  ·  2026-07-15  ·  第 ")
    set_run_font(run, size=9, color=MUTED)
    add_field(p, "PAGE")
    run = p.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_header_footer(section):
    for header in (section.header, section.even_page_header, section.first_page_header):
        header.is_linked_to_previous = False
        populate_header(header)
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        populate_footer(footer)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(run, size=9.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("用户标签物料")
    set_run_font(run, size=25, color="111827", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run("统一用户标识、奖励资产与全场景展示规范")
    set_run_font(run, size=14, color=MUTED)

    for label, value in (
        ("版本", "V1.0"),
        ("状态", "待评审"),
        ("日期", "2026-07-15"),
        ("涉及端", "iOS / Android / 服务端 / 运营后台"),
        ("协作角色", "产品、视觉设计、客户端、服务端、测试、数据、运营"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, color="111827", bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color="111827")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(11)
    set_paragraph_border_bottom(rule, color=BLUE, size=12, space=5)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], 9360, indent_dxa=120)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    set_cell_border(cell, color="B8CCE4", size=5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("核心方案")
    set_run_font(run, size=10, color=BLUE, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(
        p,
        "完整资料场景采用“等级身份行 / 资料属性行 / 荣誉标签行”；高密度场景使用固定槽位和硬上限。VIP、财富、魅力保持同一行，新奖励标签独立成行并支持最多佩戴 3 个。",
        default_size=10.5,
    )


def split_table_row(line):
    return [cell.strip().replace("<br>", "\n") for cell in line.strip().strip("|").split("|")]


def is_table_separator(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)


def choose_widths(headers, rows, total_dxa):
    n = len(headers)
    if "场景组（覆盖页面）" in headers:
        return [450, 3000, 1400, 700, 800, 3300, 2310]
    if headers[:2] == ["行", "内容"] and n == 5:
        raw = [900, 1900, 2800, 900, 2860]
        return raw
    if headers and headers[0] == "字段" and n == 2:
        return [2300, total_dxa - 2300]
    if n == 2:
        return [int(total_dxa * 0.30), total_dxa - int(total_dxa * 0.30)]
    if n == 3:
        return [int(total_dxa * 0.20), int(total_dxa * 0.36), total_dxa - int(total_dxa * 0.56)]
    lengths = []
    for idx, header in enumerate(headers):
        values = [header] + [r[idx] if idx < len(r) else "" for r in rows]
        max_len = max((len(v) for v in values), default=1)
        lengths.append(min(max(max_len, 7), 24))
    min_width = 850 if n <= 5 else 600
    remaining = total_dxa - min_width * n
    total_weight = sum(lengths)
    widths = [min_width + int(remaining * w / total_weight) for w in lengths]
    widths[-1] += total_dxa - sum(widths)
    return widths


def set_table_geometry(table, widths, total_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, landscape=False):
    total_dxa = 12960 if landscape else 9360
    widths = choose_widths(headers, rows, total_dxa)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, total_dxa, indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    font_size = 7.7 if len(headers) >= 6 else (8.3 if len(headers) == 5 else 9)

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, size=font_size, color="111827", bold=True)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 10 and "说明" not in headers[idx] else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, default_size=font_size, default_color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], 9360, indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    set_cell_border(cell, color="B8CCE4", size=4)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, default_size=10.5, default_color=DARK_BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_screenshot(doc):
    if not IMAGE_PATH.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(IMAGE_PATH), width=Inches(2.25))
    p = doc.add_paragraph(style="Caption Custom")
    p.add_run("图 1　现状示例：排队列表中的 VIP 与圆形等级标识")


def parse_markdown(doc, text, bullet_num_id, decimal_num_id):
    lines = text.splitlines()
    i = 0
    # Skip source masthead; the Word-specific masthead is already created.
    while i < len(lines) and not lines[i].startswith("## 0."):
        i += 1

    screenshot_added = False
    current_landscape = False
    landscape_scope_section6 = False
    decimal_sequence_active = False
    current_decimal_num_id = decimal_num_id
    decimal_group_len = 0
    decimal_group_pos = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip() or line.strip() == "---":
            decimal_sequence_active = False
            i += 1
            continue

        if line.startswith("#"):
            decimal_sequence_active = False
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            if title.startswith("6. 全场景展示上限与排序") and not current_landscape:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                current_landscape = True
                landscape_scope_section6 = True
            elif title.startswith("7. 统一排序、互斥与溢出算法") and current_landscape:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=False)
                current_landscape = False
                landscape_scope_section6 = False
            if level == 2:
                p = doc.add_paragraph(style="Heading 1")
            elif level == 3:
                p = doc.add_paragraph(style="Heading 2")
            else:
                p = doc.add_paragraph(style="Heading 3")
            add_inline(p, title, default_size={2: 16, 3: 13}.get(level, 12),
                       default_color={2: BLUE, 3: BLUE}.get(level, DARK_BLUE))
            i += 1
            continue

        if line.startswith("> "):
            decimal_sequence_active = False
            add_callout(doc, line[2:].strip())
            i += 1
            continue

        if i + 1 < len(lines) and line.strip().startswith("|") and is_table_separator(lines[i + 1]):
            decimal_sequence_active = False
            headers = split_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                values = split_table_row(lines[i])
                if len(values) == len(headers):
                    rows.append(values)
                i += 1
            use_landscape = len(headers) >= 6
            table_started_landscape = False
            if use_landscape and not current_landscape:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                current_landscape = True
                table_started_landscape = True
            add_table(doc, headers, rows, landscape=use_landscape)
            if use_landscape and table_started_landscape and not landscape_scope_section6:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=False)
                current_landscape = False
            continue

        bullet_match = re.match(r"^\s*-\s+(.*)$", line)
        numbered_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if bullet_match:
            decimal_sequence_active = False
            content = bullet_match.group(1)
            checked = None
            if content.startswith("[ ] "):
                checked = False
                content = content[4:]
            elif content.startswith("[x] ") or content.startswith("[X] "):
                checked = True
                content = content[4:]
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id)
            if checked is not None:
                prefix = "☒ " if checked else "☐ "
                run = p.add_run(prefix)
                set_run_font(run, size=11, color=BLUE)
            add_inline(p, content)
            i += 1
            continue
        if numbered_match:
            if not decimal_sequence_active:
                current_decimal_num_id = new_num_instance(doc, decimal_num_id)
                j = i
                decimal_group_len = 0
                while j < len(lines) and re.match(r"^\s*\d+\.\s+", lines[j].rstrip()):
                    decimal_group_len += 1
                    j += 1
                decimal_group_pos = 0
            p = doc.add_paragraph()
            apply_num(p, current_decimal_num_id)
            add_inline(p, numbered_match.group(1))
            if decimal_group_len <= 5 and decimal_group_pos < decimal_group_len - 1:
                p.paragraph_format.keep_with_next = True
            decimal_group_pos += 1
            decimal_sequence_active = True
            i += 1
            continue

        # Merge consecutive prose lines into one paragraph, preserving explicit Markdown hard breaks.
        decimal_sequence_active = False
        parts = [line.strip().rstrip("  ")]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (not nxt.strip() or nxt.startswith("#") or nxt.startswith("> ") or
                    re.match(r"^\s*-\s+", nxt) or re.match(r"^\s*\d+\.\s+", nxt) or
                    nxt.strip().startswith("|") or nxt.strip() == "---"):
                break
            parts.append(nxt.strip().rstrip("  "))
            i += 1
        text_value = " ".join(parts)
        p = doc.add_paragraph()
        add_inline(p, text_value)

        if (not screenshot_added and "附件截图中的“排队列表”" in text_value):
            add_screenshot(doc)
            screenshot_added = True


def audit_document(doc):
    for section in doc.sections:
        assert abs(section.left_margin - Inches(1)) < 100
        assert abs(section.right_margin - Inches(1)) < 100
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa"
        assert table._tbl.tblGrid is not None
        for row in table.rows:
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tc_w is not None and tc_w.get(qn("w:type")) == "dxa"


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    configure_section(doc.sections[0], landscape=False)
    add_header_footer(doc.sections[0])
    bullet_num_id, decimal_num_id = add_numbering(doc)

    props = doc.core_properties
    props.title = "用户标签物料产品需求文档"
    props.subject = "统一用户标识、奖励资产与全场景展示规范"
    props.author = "产品团队"
    props.keywords = "PRD, 用户标签, 徽章, 荣誉标签, 展示规范"
    props.comments = "V1.0 待评审"

    add_masthead(doc)
    parse_markdown(doc, MD_PATH.read_text(encoding="utf-8"), bullet_num_id, decimal_num_id)
    audit_document(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
