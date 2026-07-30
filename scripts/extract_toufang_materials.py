from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return {
        "file": str(path),
        "type": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def cell_value(cell):
    value = cell.value
    if value is None:
        return None
    if cell.data_type == "f":
        return f"={value}"
    return value


def extract_xlsx(path: Path) -> dict:
    workbook = load_workbook(path, data_only=False)
    sheets = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows():
            values = [cell_value(cell) for cell in row]
            if any(v not in (None, "") for v in values):
                while values and values[-1] in (None, ""):
                    values.pop()
                rows.append(values)
        sheets.append({"name": sheet.title, "rows": rows})
    return {
        "file": str(path),
        "type": "xlsx",
        "sheets": sheets,
    }


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: extract_toufang_materials.py <source_dir> <output_json>", file=sys.stderr)
        return 2

    source_dir = Path(sys.argv[1])
    output_json = Path(sys.argv[2])
    items = []
    for path in sorted(source_dir.iterdir()):
        if path.suffix.lower() == ".docx":
            items.append(extract_docx(path))
        elif path.suffix.lower() == ".xlsx":
            items.append(extract_xlsx(path))

    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_json.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")
    print(output_json)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
