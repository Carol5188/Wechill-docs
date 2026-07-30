from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


EAST_ASIA_FONT = "PingFang SC"
LATIN_FONT = "Aptos"


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clean_inline(text: str) -> str:
    return text.replace("`", "").strip()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [clean_inline(cell) for cell in raw.strip("|").split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)
        if not is_separator:
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            cell = table.cell(row_index, col_index)
            text = row[col_index] if col_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_run_font(run, 8, bold=(row_index == 0), color="1F2937")
            if row_index == 0:
                set_cell_shading(cell, "E5E7EB")

    doc.add_paragraph()


def add_paragraph_with_inline_styles(doc: Document, text: str, style: str | None = None, size: int = 10) -> None:
    paragraph = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part.strip("`"))
            set_run_font(run, size, color="374151")
            run.font.name = "Menlo"
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, color="111827")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    for style_name, size in [
        ("Normal", 10),
        ("List Bullet", 10),
        ("List Number", 10),
        ("Quote", 9),
    ]:
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)

    for style_name, size, color in [
        ("Title", 21, "111827"),
        ("Heading 1", 15, "0F172A"),
        ("Heading 2", 12, "1F2937"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    i = 0
    title_added = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(clean_inline(stripped[2:]))
            set_run_font(run, 21, bold=True, color="111827")
            title_added = True
            i += 1
            continue

        if stripped.startswith("## "):
            if title_added:
                doc.add_section(WD_SECTION_START.CONTINUOUS)
            paragraph = doc.add_paragraph(style="Heading 1")
            run = paragraph.add_run(clean_inline(stripped[3:]))
            set_run_font(run, 15, bold=True, color="0F172A")
            i += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            run = paragraph.add_run(clean_inline(stripped[4:]))
            set_run_font(run, 12, bold=True, color="1F2937")
            i += 1
            continue

        if stripped.startswith("#### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            run = paragraph.add_run(clean_inline(stripped[5:]))
            set_run_font(run, 11, bold=True, color="374151")
            i += 1
            continue

        if stripped.startswith(">"):
            paragraph = doc.add_paragraph(style="Quote")
            run = paragraph.add_run(clean_inline(stripped.lstrip("> ")))
            set_run_font(run, 9, color="4B5563")
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if numbered:
            add_paragraph_with_inline_styles(doc, stripped)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet:
            add_paragraph_with_inline_styles(doc, bullet.group(1), style="List Bullet")
            i += 1
            continue

        add_paragraph_with_inline_styles(doc, stripped)
        i += 1

    doc.save(docx_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate_gcc_toufang_docx.py input.md output.docx")
    markdown_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
