from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / "Lucky Fruit Party活动说明.docx"
LONG_IMG = ROOT / "Lucky Fruit Party-活动长图-单张.png"
PREVIEW_IMG = ROOT / "00-当前页面.png"

PURPLE = "5B21B6"
MAGENTA = "B000D4"
GOLD = "D4A72C"
INK = "24133B"
MUTED = "6B5B76"
PALE = "F8F1FF"
LAVENDER = "EEE4FF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Arial Unicode MS", size=10.5, color=INK, bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_para(paragraph, space_before=0, space_after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line


def add_text(p, text, size=10.5, color=INK, bold=False):
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return r


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para(p, space_before=8 if level == 1 else 4, space_after=5, line=1.0)
    add_text(p, text, size=16 if level == 1 else 12, color=PURPLE if level == 1 else MAGENTA, bold=True)
    return p


def add_bullet(doc, text, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, space_after=3, line=1.12)
    add_text(p, text, size=10.2, color=color)
    return p


def style_table(table, header_fill=PURPLE):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                set_para(p, space_after=0, line=1.08)
                for r in p.runs:
                    set_run_font(r, size=9.3, color="FFFFFF" if ri == 0 else INK, bold=ri == 0)
            if ri == 0:
                set_cell_shading(cell, header_fill)
            elif ri % 2 == 0:
                set_cell_shading(cell, PALE)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, space_after=0, line=1.0)
    add_text(p, "Lucky Fruit Party · 活动页面整理 · 2026-07-27", size=8.5, color=MUTED)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.66)
section.right_margin = Inches(0.66)
add_footer(section)

styles = doc.styles
styles["Normal"].font.name = "Arial Unicode MS"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
styles["Normal"].font.size = Pt(10.5)

# Cover / executive summary
p = doc.add_paragraph()
set_para(p, space_after=2, line=1.0)
add_text(p, "LUCKY FRUIT PARTY", size=26, color=PURPLE, bold=True)
p = doc.add_paragraph()
set_para(p, space_after=10, line=1.0)
add_text(p, "活动详情、奖励与玩法速览", size=18, color=MAGENTA, bold=True)

p = doc.add_paragraph()
set_para(p, space_after=12, line=1.15)
add_text(p, "来源：通过 scrcpy 查看手机当前活动页面并整理。截图时页面用户为 Haris06，活动积分显示为 0。", size=9.5, color=MUTED)

add_heading(doc, "一、活动概览", 1)
t = doc.add_table(rows=1, cols=2)
overview = [
    ("活动名称", "Lucky Fruit Party"),
    ("活动时间", "GMT+2：Jul.1 00:00 – Jul.31 23:59"),
    ("活动入口页签", "Daily Task；Weekly Ranking"),
]
for _ in overview:
    t.add_row()
for i, (a, b) in enumerate(overview, 1):
    t.cell(i, 0).text = a
    t.cell(i, 1).text = b
style_table(t)

add_heading(doc, "二、玩法一句话", 1)
callout = doc.add_table(rows=1, cols=1)
cell = callout.cell(0, 0)
set_cell_shading(cell, "F4D9FF")
set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para(p, space_after=0, line=1.2)
add_text(p, "在 “Lucky Fruit” 游戏中每赢得 1 个 gold coin，即获得 1 个活动积分。", size=13, color=PURPLE, bold=True)

add_bullet(doc, "Daily Task：累计活动积分，逐档解锁 50K、250K、500K、2.5M、5M 目标奖励。")
add_bullet(doc, "Weekly Ranking：按周榜积分排名，页面展示 Top1 / Top2 / Top3 奖励。")
add_bullet(doc, "奖励状态：本次截图中各档均显示 Unachieved；图标奖励的具体商品名未在页面文字中显示，文档按图标和时长记录。")

doc.add_page_break()

# Daily task
add_heading(doc, "三、Daily Task：累计积分奖励", 1)
p = doc.add_paragraph()
set_para(p, space_after=8, line=1.15)
add_text(p, "页面将累计积分设计为五个里程碑。每档奖励下方标注 x 1 days，表示页面展示的 1 天时效。", size=10.3, color=INK)

t = doc.add_table(rows=1, cols=3)
headers = ["目标积分", "页面展示奖励", "截图状态"]
for i, h in enumerate(headers):
    t.cell(0, i).text = h
rows = [
    ("50K", "1 个猫咪主题装扮/头像框类图标；x 1 days", "Unachieved"),
    ("250K", "1 个童趣人物主题装扮/头像框类图标；x 1 days", "Unachieved"),
    ("500K", "2 个图标奖励（龙/果实主题）；各 x 1 days", "Unachieved"),
    ("2.5M", "2 个图标奖励（情侣/果实主题）；各 x 1 days", "Unachieved"),
    ("5M", "3 个图标奖励（科幻/果实/黄金载具主题）；各 x 1 days", "Unachieved"),
]
for row in rows:
    cells = t.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
style_table(t, header_fill=MAGENTA)

add_heading(doc, "解读", 2)
add_bullet(doc, "积分门槛是累计值，不是单次赢取值；页面没有显示每日重置提示。")
add_bullet(doc, "奖励图标以视觉资产呈现，页面未给出文字商品名；若用于 PRD 或配置表，建议再从后台配置核对正式名称。")
add_bullet(doc, "目标栏上方的 50K / 250K / 500K / 2.5M / 5M 与下方 Target 档位一一对应。")

doc.add_page_break()

# Weekly ranking
add_heading(doc, "四、Weekly Ranking：周榜奖励", 1)
p = doc.add_paragraph()
set_para(p, space_after=8, line=1.15)
add_text(p, "页面显示按周排名的前三名奖励。当前用户 Haris06 在截图时为 0 分，榜单中显示为未上榜/占位状态。", size=10.3, color=INK)

t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["名次", "数值奖励", "时效图标奖励"]):
    t.cell(0, i).text = h
weekly = [
    ("Top 1", "+8000；+50000 EXP ×3", "3 个图标，页面均显示 x 7 days"),
    ("Top 2", "+6000；+50000 EXP ×2", "3 个图标，页面均显示 x 5 days"),
    ("Top 3", "+3500；+50000 EXP ×1", "3 个图标，页面均显示 x 3 days"),
]
for row in weekly:
    cells = t.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
style_table(t, header_fill=PURPLE)

add_heading(doc, "五、执行流程", 1)
for text in [
    "进入活动页，确认活动时间与当前积分。",
    "进入 Lucky Fruit 游戏并进行游戏；每赢得 1 个 gold coin，活动积分 +1。",
    "回到活动页查看 Daily Task 进度，达到目标后领取/解锁对应奖励（页面按钮状态从 Unachieved 变化）。",
    "需要竞争周榜时，切换 Weekly Ranking，关注周榜名次及奖励时效。",
]:
    add_bullet(doc, text)

add_heading(doc, "六、页面观察备注", 1)
for text in [
    "活动页面使用老虎机、水果、金币和舞台幕布作为统一视觉主题。",
    "活动倒计时位于活动时间下方；截图时显示约 4 day 14 hour 30 minute，具体数值会随时间变化。",
    "奖励图标未在页面中显示文字名称，本文避免将图标强行对应到未确认的正式商品名。",
]:
    add_bullet(doc, text, color=MUTED)

doc.add_page_break()

# Appendix with the single long screenshot.
add_heading(doc, "附录：活动页面单张长图", 1)
p = doc.add_paragraph()
set_para(p, space_after=7, line=1.15)
add_text(p, "以下为从手机页面连续滚动后合成的单张长图预览；完整原图另存为：", size=10.2, color=INK)
add_text(p, "Lucky Fruit Party-活动长图-单张.png", size=10.2, color=PURPLE, bold=True)

layout = doc.add_table(rows=1, cols=2)
layout.alignment = WD_TABLE_ALIGNMENT.CENTER
layout.autofit = False
layout.columns[0].width = Inches(4.55)
layout.columns[1].width = Inches(2.65)
left, right = layout.rows[0].cells
set_cell_margins(left, top=120, start=140, bottom=120, end=160)
set_cell_margins(right, top=80, start=80, bottom=80, end=80)
set_cell_shading(left, PALE)
set_cell_shading(right, "FFFFFF")

lp = left.paragraphs[0]
set_para(lp, space_after=5, line=1.12)
add_text(lp, "长图包含：", size=11, color=PURPLE, bold=True)
for item in [
    "活动首页与活动时间",
    "Daily Task / Weekly Ranking 页签",
    "50K、250K、500K、2.5M、5M 五档目标",
    "各档奖励图标与 Unachieved 状态",
]:
    p = left.add_paragraph(style="List Bullet")
    set_para(p, space_after=4, line=1.12)
    add_text(p, item, size=10.2, color=INK)

rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para(rp, space_after=3, line=1.0)
run = rp.add_run()
run.add_picture(str(PREVIEW_IMG), width=Inches(2.20))
cap = right.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para(cap, space_before=2, space_after=0, line=1.0)
add_text(cap, "首页截图预览；完整单张长图另存为 PNG", size=8.5, color=MUTED)

doc.save(OUT)
print(OUT)
